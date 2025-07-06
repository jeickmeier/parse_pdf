import asyncio
from pathlib import Path

import docx
import pytest

from doc_parser.config import AppConfig as Settings
from doc_parser.parsers.docx.parser import DocxParser


@pytest.fixture()
def make_sample_docx(tmp_path):
    """Return path to a dynamically generated small DOCX file."""

    def _factory() -> Path:
        file_path = tmp_path / "sample.docx"
        document = docx.Document()

        # Heading and paragraph
        document.add_heading("Title", level=1)
        document.add_paragraph("This is a paragraph.")

        # Table 2x2
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "H1"
        table.cell(0, 1).text = "H2"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"

        document.save(file_path)
        return file_path

    return _factory


@pytest.mark.asyncio
async def test_docx_parser_markdown(make_sample_docx):
    docx_path = make_sample_docx()

    settings = Settings(
        output_format="markdown",
        parser_settings={"docx": {"extract_images": False}},
    )
    parser = DocxParser(settings)

    result = await parser.parse(docx_path)
    assert "# Title" in result.content
    assert "This is a paragraph." in result.content
    # Table markdown header row
    assert "| H1 | H2 |" in result.content

    # metadata checks
    assert result.metadata["paragraphs"] >= 2
    assert result.metadata["tables"] == 1
    assert result.format == "markdown"


@pytest.mark.asyncio
async def test_docx_parser_json(make_sample_docx):
    docx_path = make_sample_docx()

    settings = Settings(output_format="json")
    parser = DocxParser(settings)

    result = await parser.parse(docx_path)
    assert "\"paragraphs\"" in result.content
    assert "\"tables\"" in result.content


def test_docx_validate_input_neg(tmp_path):
    fake_path = tmp_path / "not_docx.txt"
    fake_path.write_text("x")

    parser = DocxParser(Settings())
    assert asyncio.run(parser.validate_input(fake_path)) is False 